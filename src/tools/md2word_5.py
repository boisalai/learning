from typing import Optional, Dict, List, Tuple
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import os
import subprocess
import shutil

class DocumentStyle(Enum):
    REPORT = "report"
    NOTE = "note"
    LETTER = "letter"
    MEMO = "memo"

class DocumentConfig:
    def __init__(
        self,
        style: DocumentStyle = DocumentStyle.REPORT,
        title: str = "",
        author: str = "",
        date: str = "",
        heading_colors: Dict[int, Tuple[int, int, int]] = None,
        footer_text: Dict[str, str] = None,
        font_name: str = "Arial",
        base_font_size: int = 12,
        margins: Tuple[float, float, float, float] = (2, 2, 2, 2),  # top, right, bottom, left in cm
        line_spacing: float = 1.0
    ):
        self.style = style
        self.title = title
        self.author = author
        self.date = date
        self.heading_colors = heading_colors or {
            1: (37, 150, 190),  # Default blue color for all headings
            2: (37, 150, 190),
            3: (37, 150, 190)
        }
        self.footer_text = footer_text or {
            "odd": "Page",
            "even": "Page"
        }
        self.font_name = font_name
        self.base_font_size = base_font_size
        self.margins = margins
        self.line_spacing = line_spacing

    @classmethod
    def create_report_style(cls, **kwargs):
        """Style préconfiguré pour les rapports"""
        default_config = {
            'style': DocumentStyle.REPORT,
            'heading_colors': {
                1: (37, 150, 190),
                2: (37, 150, 190),
                3: (37, 150, 190)
            },
            'footer_text': {
                "odd": "Rapport | Page",
                "even": "Page | Rapport"
            }
        }
        default_config.update(kwargs)
        return cls(**default_config)

    @classmethod
    def create_note_style(cls, **kwargs):
        """Style préconfiguré pour les notes"""
        default_config = {
            'style': DocumentStyle.NOTE,
            'heading_colors': {
                1: (70, 70, 70),
                2: (100, 100, 100),
                3: (130, 130, 130)
            },
            'footer_text': {
                "odd": "Note interne | Page",
                "even": "Page | Note interne"
            },
            'margins': (1.5, 1.5, 1.5, 1.5)
        }
        default_config.update(kwargs)
        return cls(**default_config)

class MarkdownToDocxConverter:
    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Vérifie que Pandoc est installé"""
        if not shutil.which('pandoc'):
            raise RuntimeError(
                "Pandoc n'est pas installé. "
                "Sur macOS, installez-le avec: brew install pandoc"
            )

    def convert(self, 
                input_file: str,
                output_file: str,
                working_dir: Optional[str] = None,
                extra_args: Optional[List[str]] = None) -> None:
        """
        Convertit un fichier Markdown en Word avec la configuration spécifiée
        
        Args:
            input_file: Nom du fichier markdown d'entrée
            output_file: Nom du fichier Word de sortie
            working_dir: Répertoire de travail (contenant le markdown et où sera créé le docx)
            extra_args: Arguments supplémentaires pour pandoc
        """
        # Définir le répertoire de travail
        work_dir = Path(working_dir).resolve() if working_dir else Path.cwd()
        input_path = work_dir / input_file
        output_path = work_dir / output_file
        
        # Vérifier l'existence du fichier d'entrée
        if not input_path.exists():
            raise FileNotFoundError(f"Le fichier d'entrée n'existe pas: {input_path}")
        
        # Créer le répertoire de sortie si nécessaire
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Vérifier et créer le répertoire des images si nécessaire
        img_dir = work_dir / "img"
        if not img_dir.exists():
            img_dir.mkdir(parents=True)
            print(f"Répertoire d'images créé: {img_dir}")
        
        # Lire le contenu du fichier markdown pour vérifier les images
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Ajuster les chemins des images dans le contenu markdown
        content = self._process_image_paths(content, img_dir)
        
        # Créer un fichier markdown temporaire avec les chemins ajustés
        temp_md = work_dir / f"temp_{input_path.name}"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Commande pandoc
        cmd = [
            'pandoc',
            str(temp_md),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999',
            '--resource-path', str(img_dir),  # Correction ici : '--resource-path' en minuscules
            '-M', f'title={self.config.title}',
            '-M', f'author={self.config.author}',
            '-M', f'date={self.config.date}'
        ]
        
        if extra_args:
            cmd.extend(extra_args)
        
        try:
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            self._post_process_document(output_path)
            print(f"Conversion réussie! Fichier sauvegardé: {output_path}")
            
        except subprocess.CalledProcessError as e:
            print(f"Erreur lors de la conversion: {e.stderr}")
            raise
            
        finally:
            # Nettoyer le fichier temporaire
            if temp_md.exists():
                temp_md.unlink()

    def _process_image_paths(self, content: str, img_dir: Path) -> str:
        """Traite les chemins des images dans le contenu markdown"""
        import re
        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        
        def process_image_match(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            img_name = Path(img_path).name
            img_file = img_dir / img_name
            
            if not img_file.exists():
                print(f"Attention: Image non trouvée: {img_file}")
                return match.group(0)
                
            return f'![{alt_text}](img/{img_name})'
            
        return re.sub(img_pattern, process_image_match, content)

    def _post_process_document(self, doc_path: Path) -> None:
        """
        Post-traitement du document Word pour appliquer tous les styles et la mise en forme.
        """
        doc = Document(doc_path)

        # Application des styles et configurations globales
        self._apply_global_styles(doc)

        # Traitement des paragraphes spéciaux (listes, titres, code, etc.)
        self._process_paragraphs(doc, doc_path)
        
        # Traitement des tableaux
        self._process_tables(doc)
        
        # Forcer le style du titre (premier paragraphe)
        if doc.paragraphs:
            title_para = doc.paragraphs[0]
            title_para.clear()
            title_para.style = doc.styles['Title']
            run = title_para.add_run(self.config.title)
            run.font.name = self.config.font_name
            run.font.size = Pt(24)
            run.font.bold = True

        # Sauvegarder les modifications
        doc.save(doc_path)

    def _apply_global_styles(self, doc: Document) -> None:
        """Applique les styles et configurations globaux au document"""
        try:
            title_style = doc.styles['Title']
            title_style.font.name = self.config.font_name
            title_style.font.size = Pt(24)
            title_style.font.bold = True
        except KeyError:
            print("Style 'Title' non trouvé")

        style_normal = doc.styles['Normal']
        style_normal.font.name = self.config.font_name
        style_normal.font.size = Pt(self.config.base_font_size)
        style_normal.paragraph_format.line_spacing = self.config.line_spacing

        try:
            style_default = doc.styles['Default Paragraph Font']
            style_default.font.name = self.config.font_name
            style_default.font.size = Pt(self.config.base_font_size)
        except KeyError:
            print("Style 'Default Paragraph Font' non trouvé")

        for style_name in ['Body Text', 'Normal', 'Plain Text']:
            try:
                style = doc.styles[style_name]
                style.font.name = self.config.font_name
                style.font.size = Pt(self.config.base_font_size)
            except KeyError:
                continue

        self._adjust_toc_styles(doc)
        self._adjust_heading_number_spacing(doc)
        self._configure_section_properties(doc)
        self._setup_footers(doc)
        self._process_footnotes(doc)
        
        for section in doc.sections:
            section.top_margin = Cm(self.config.margins[0])
            section.bottom_margin = Cm(self.config.margins[1])
            section.left_margin = Cm(self.config.margins[2])
            section.right_margin = Cm(self.config.margins[3])

    def _process_paragraphs(self, doc: Document, doc_path: Path) -> None:
        """Traite tous les paragraphes du document"""
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):
                for run in para.runs:
                    run.font.name = self.config.font_name
                    run.font.size = Pt(self.config.base_font_size)

            if self._is_list_paragraph(para.text):
                self._remove_list_spacing(para)
            
            if para.style.name.startswith('Heading'):
                self._apply_heading_style(para)
            elif 'CodeBlock' in para.style.name:
                self._apply_code_style(para)
            elif "Logo Markdown" in para.text:
                self._process_logo_image(para, doc_path)

    def _is_list_paragraph(self, text: str) -> bool:
        """Détermine si un paragraphe est une liste"""
        return ('•' in text or 
                text.strip().startswith('o') or 
                text.strip().startswith('§') or
                any(str(i)+'.' in text for i in range(10)))

    def _apply_heading_style(self, para) -> None:
        """Applique le style aux titres"""
        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        for run in para.runs:
            run.font.name = self.config.font_name
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.config.heading_colors.get(level, (37, 150, 190)))

    def _apply_code_style(self, para) -> None:
        """Applique le style aux blocs de code"""
        for run in para.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        para.paragraph_format.left_indent = Cm(1)

    def _process_logo_image(self, para, doc_path: Path) -> None:
        """Traite l'image du logo Markdown"""
        image_path = doc_path.parent / "img" / "markdown.png"
        if image_path.exists():
            para.text = ""
            run = para.add_run()
            run.add_picture(str(image_path), width=Inches(3))
        else:
            print(f"Image non trouvée: {image_path}")

    def _process_tables(self, doc: Document) -> None:
        """Traite tous les tableaux du document"""
        for table in doc.tables:
            self._process_table_header(table)
            for i, row in enumerate(table.rows):
                if i > 0:  
                    for cell in row.cells:
                        self._process_table_cell(cell, is_header=False)

    def _process_table_header(self, table) -> None:
        """Traite la première ligne du tableau (en-tête)"""
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                self._process_table_cell(cell, is_header=True)

    def _process_table_cell(self, cell: _Cell, is_header: bool = False) -> None:
        """Traite une cellule de tableau"""
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            table_font_size = self.config.base_font_size - 1
            for run in para.runs:
                run.font.name = self.config.font_name
                run.font.size = Pt(table_font_size)
                if is_header:
                    run.font.bold = True
        
            if not para.text.strip():
                run = para.add_run(" ")
                run.font.name = self.config.font_name
                run.font.size = Pt(table_font_size)
                if is_header:
                    run.font.bold = True
                    
        self._set_cell_borders(cell)

    def _set_cell_borders(self, cell: _Cell) -> None:
        """Ajoute des bordures à une cellule"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '4')
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), 'auto')
            tcBorders.append(border_el)

    def _remove_list_spacing(self, paragraph) -> None:
        """Supprime les espacements d'un paragraphe au niveau XML"""
        pPr = paragraph._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        old_spacing = pPr.find(qn('w:spacing'))
        if old_spacing is not None:
            pPr.remove(old_spacing)
        pPr.append(spacing)

    def _configure_section_properties(self, doc) -> None:
        """Configure les propriétés de section pour les en-têtes et pieds de page"""
        for section in doc.sections:
            sectPr = section._sectPr
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

    def _setup_footers(self, doc) -> None:
        """Configuration des pieds de page avec texte personnalisé"""
        for section in doc.sections:
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True

            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()
            p_odd.paragraph_format.space_before = Pt(12)
            run_odd = p_odd.add_run(f"{self.config.footer_text['odd']} | ")
            run_odd.font.name = self.config.font_name
            run_odd.font.size = Pt(10)
            self._add_page_number(p_odd)
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()
            p_even.paragraph_format.space_before = Pt(12)
            self._add_page_number(p_even)
            run_even = p_even.add_run(f" | {self.config.footer_text['even']}")
            run_even.font.name = self.config.font_name
            run_even.font.size = Pt(10)
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_page_number(self, paragraph) -> None:
        """Ajoute un numéro de page au paragraphe"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _process_footnotes(self, doc) -> None:
        """Configure les notes de bas de page"""
        try:
            styles = doc.styles
            style = styles['Footnote Text']
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            ref_style = styles['Footnote Reference']
            ref_style.font.name = 'Arial'
            ref_style.font.size = Pt(10)
            
            if hasattr(doc.part, '_footnotes_part') and doc.part._footnotes_part:
                footnotes = doc.part._footnotes_part.element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
                )
                for footnote in footnotes:
                    pPr = footnote.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr'
                    )
                    if pPr is None:
                        pPr = parse_xml(
                            r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        )
                        footnote.insert(0, pPr)
                    spacing = parse_xml(
                        r'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        r'w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
                    )
                    old_spacing = pPr.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing'
                    )
                    if old_spacing is not None:
                        pPr.remove(old_spacing)
                    pPr.append(spacing)

        except Exception as e:
            print(f"Avertissement lors du traitement des notes de bas de page : {str(e)}")

    def _adjust_toc_styles(self, doc) -> None:
        """Supprime le gras des paragraphes de la table des matières"""
        in_toc = False
        for paragraph in doc.paragraphs:
            if paragraph.text in ["Table des matières", "Table of Contents"]:
                in_toc = True
                continue
            if in_toc and paragraph.style.name.startswith('Heading 1'):
                in_toc = False
                continue
            if in_toc:
                for run in paragraph.runs:
                    run.font.bold = False
                    if hasattr(run, '_element') and run._element.rPr is not None:
                        rPr = run._element.rPr
                        for bold in rPr.findall(qn('w:b')):
                            rPr.remove(bold)
                        bold = OxmlElement('w:b')
                        bold.set(qn('w:val'), '0')
                        rPr.append(bold)

    def _adjust_heading_number_spacing(self, doc) -> None:
        """Ajuste l'espacement entre la numérotation et le texte des titres"""
        for style in doc.styles:
            if style.name.startswith('Heading'):
                if hasattr(style, '_element') and style._element.pPr is not None:
                    pPr = style._element.pPr
                    numPr = pPr.get_or_add_numPr()
                    tabs = pPr.get_or_add_tabs()
                    for tab in tabs.findall(qn('w:tab')):
                        tabs.remove(tab)
                    tab = OxmlElement('w:tab')
                    tab.set(qn('w:val'), 'left')
                    tab.set(qn('w:pos'), '960')
                    tabs.append(tab)
                    ind = pPr.get_or_add_ind()
                    ind.set(qn('w:left'), '0')
                    ind.set(qn('w:firstLine'), '0')
                    ind.set(qn('w:hanging'), '0')

def create_doc1():
    try:
        converter = MarkdownToDocxConverter(
            config=DocumentConfig.create_report_style(
                title="Mon Rapport",
                author="John Doe",
                date="2024-01-10"
            )
        )
        converter.convert(
            input_file="input.md",
            output_file="output1.docx",
            working_dir="/Users/alain/Downloads",
            extra_args=['--toc', '--number-sections']
        )
    except Exception as e:
        print(f"Erreur: {str(e)}")

def create_doc2():
    try:
        config = DocumentConfig.create_report_style(
            title="Mon Rapport",
            author="John Doe",
            date="2024-01-10",
            font_name="Arial", 
            base_font_size=12,  
            heading_colors={
                1: (0, 0, 0),
                2: (0, 0, 0),
                3: (0, 0, 0)
            },
            footer_text={
                "odd": "Mon Entreprise",
                "even": "Rapport confidentiel"
            }
        )
        converter = MarkdownToDocxConverter(config)
        converter.convert(
            input_file="input.md",
            output_file="output3.docx",
            working_dir="/Users/alain/Workspace/GitHub/Learning/src/tools",
            extra_args=['--toc', '--number-sections']
        )
    except Exception as e:
        print(f"Erreur: {str(e)}")

if __name__ == "__main__":
    # create_doc1()
    create_doc2()
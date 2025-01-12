from typing import Optional, Dict, List, Tuple
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import os
import subprocess
import shutil
import re


class PaperSize(Enum):
    LETTER = "letter"  # 8.5 x 11 inches (215.9 x 279.4 mm)
    LEGAL = "legal"    # 8.5 x 14 inches (215.9 x 355.6 mm)
    A4 = "a4"         # 8.27 x 11.69 inches (210 x 297 mm)

class DocumentStyle(Enum):
    REPORT = "report"
    NOTE = "note"
    LETTER = "letter"
    MEMO = "memo"

class DocumentConfig:
    def __init__(
        self,
        style: DocumentStyle = DocumentStyle.REPORT,
        paper_size: PaperSize = PaperSize.LETTER,
        author: str = "",
        date: str = "",
        heading_colors: Dict[int, Tuple[int, int, int]] = None,
        footer_text: Dict[str, str] = None,
        font_name: str = "Arial",
        base_font_size: int = 12,
        margins: Tuple[float, float, float, float] = (2, 2, 2, 2),  # top, right, bottom, left in cm
        line_spacing: float = 1.0,
        generate_toc: bool = True
    ):
        self.style = style
        self.paper_size = paper_size
        self.author = author
        self.date = date
        self.heading_colors = heading_colors or {
            1: (37, 150, 190),  # Default blue color for all headings
            2: (37, 150, 190),
            3: (37, 150, 190)
        }
        self.footer_text = footer_text or {
            "odd": "Page",
            "even": "Page"
        }
        self.font_name = font_name
        self.base_font_size = base_font_size
        self.margins = margins
        self.line_spacing = line_spacing
        self.generate_toc = generate_toc

    @classmethod
    def create_report_style(cls, **kwargs):
        """Preconfigured style for reports"""
        default_config = {
            'style': DocumentStyle.REPORT,
            'paper_size': PaperSize.LETTER,
            'heading_colors': {
                1: (37, 150, 190),
                2: (37, 150, 190),
                3: (37, 150, 190)
            },
            'footer_text': {
                "odd": "Right text | Page",
                "even": "Page | Left text"
            }
        }
        default_config.update(kwargs)
        return cls(**default_config)

    @classmethod
    def create_note_style(cls, **kwargs):
        """Preconfigured style for internal notes"""
        default_config = {
            'style': DocumentStyle.NOTE,
            'paper_size': PaperSize.LEGAL,
            'heading_colors': {
                1: (70, 70, 70),
                2: (100, 100, 100),
                3: (130, 130, 130)
            },
            'footer_text': {
                "odd": "Internal Note | Page",
                "even": "Page | Internal Note"
            },
            'margins': (1.5, 1.5, 1.5, 1.5)
        }
        default_config.update(kwargs)
        return cls(**default_config)

class MarkdownToDocxConverter:
    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if Pandoc is installed"""
        if not shutil.which('pandoc'):
            raise RuntimeError(
                "Pandoc is not installed. "
                "On macOS, install it with: brew install pandoc"
            )

    def _extract_title_from_markdown(self, content: str) -> str:
        """Extract the main title (h1) from markdown content"""
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
        return "Untitled Document"

    def _extract_image_references(self, content: str) -> List[dict]:
        """Extract image references from markdown content"""
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        matches = re.findall(image_pattern, content)
        
        image_refs = []
        for alt_text, path in matches:
            path = path.strip()
            if path.startswith('img/'):
                path = path[4:]
                
            image_refs.append({
                'alt_text': alt_text,
                'path': path,
                'original_markdown': f'![{alt_text}]({path})'
            })
        
        print(f"Found {len(image_refs)} image references")
        return image_refs

    def _remove_image_references(self, content: str) -> str:
        """Remove image references from markdown content"""
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        return re.sub(image_pattern, r'[IMAGE_PLACEHOLDER]', content)

    def convert(self, 
                input_file: str,
                output_file: str,
                working_dir: Optional[str] = None,
                extra_args: Optional[List[str]] = None) -> None:
        """Convert Markdown to Word"""
        # Set working directory
        work_dir = Path(working_dir).resolve() if working_dir else Path.cwd()
        input_path = work_dir / input_file
        output_path = work_dir / output_file
        
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Create image directory if needed
        img_dir = work_dir / "img"
        if not img_dir.exists():
            img_dir.mkdir(parents=True)
            print(f"Created image directory: {img_dir}")
        
        # Read markdown content
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract title and image references
        document_title = self._extract_title_from_markdown(content)
        image_refs = self._extract_image_references(content)
        
        # Create temporary markdown without images
        temp_content = self._remove_image_references(content)
        temp_md = work_dir / f"temp_{input_path.name}"
        
        try:
            # Write temporary markdown
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(temp_content)
            
            # Run pandoc conversion
            self._run_pandoc_conversion(temp_md, output_path, document_title)
            
            # Post-process the document
            self._post_process_document(output_path, document_title)
            
            # Insert images
            self._insert_images(output_path, image_refs, work_dir)
            
            print(f"Conversion successful! File saved: {output_path}")
            
        finally:
            # Cleanup
            if temp_md.exists():
                temp_md.unlink()

    def _run_pandoc_conversion(self, input_path: Path, output_path: Path, title: str) -> None:
        """Run pandoc conversion"""
        cmd = [
            'pandoc',
            str(input_path),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999',
            '-M', f'title={title}',
            '-M', f'author={self.config.author}',
            '-M', f'date={self.config.date}'
        ]
        
        if self.config.generate_toc:
            cmd.extend(['--toc', '--number-sections'])
            
        try:
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            print("Pandoc conversion completed successfully")
        except subprocess.CalledProcessError as e:
            print(f"Pandoc conversion failed: {e.stderr}")
            raise

    def _setup_footers(self, doc: Document) -> None:
        """Configure footers with custom text and page numbers for odd and even pages"""
        for section in doc.sections:
            # Configure section for different odd/even pages
            sectPr = section._sectPr
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

            # Configure headers/footers
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True

            # Configure odd page footer (right pages)
            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_odd.paragraph_format.space_before = Pt(12)
            
            # Add text and page number for odd pages
            run_odd = p_odd.add_run(f"{self.config.footer_text['odd']} | ")
            run_odd.font.name = self.config.font_name
            run_odd.font.size = Pt(10)
            self._add_page_number(p_odd)

            # Configure even page footer (left pages)
            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_even.paragraph_format.space_before = Pt(12)
            
            # Add page number and text for even pages
            self._add_page_number(p_even)
            run_even = p_even.add_run(f" | {self.config.footer_text['even']}")
            run_even.font.name = self.config.font_name
            run_even.font.size = Pt(10)

    def _add_page_number(self, paragraph) -> None:
        """Add a page number field to a paragraph"""
        run = paragraph.add_run()
        
        # Begin field character
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Instruction text
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        # End field character
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def _process_footnotes(self, doc: Document) -> None:
        """Configure footnotes with proper formatting"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        try:
            # Configure footnote text style
            style = doc.styles['Footnote Text']
            style.font.name = self.config.font_name
            style.font.size = Pt(10)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            
            # Configure footnote reference style
            ref_style = doc.styles['Footnote Reference']
            ref_style.font.name = self.config.font_name
            ref_style.font.size = Pt(10)
            ref_style.font.superscript = True

            # Process each footnote if they exist
            if hasattr(doc, '_part') and hasattr(doc._part, '_footnotes_part') and doc._part._footnotes_part:
                footnotes = doc._part._footnotes_part.element.xpath('//w:footnote')
                for footnote in footnotes:
                    # Process each paragraph in the footnote
                    for p in footnote.xpath('.//w:p'):
                        pPr = p.get_or_add_pPr()
                        
                        # Set paragraph properties
                        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
                        existing_spacing = pPr.find(qn('w:spacing'))
                        if existing_spacing is not None:
                            pPr.remove(existing_spacing)
                        pPr.append(spacing)
                        
                        # Process runs in the footnote
                        for r in p.xpath('.//w:r'):
                            # Set run properties if needed
                            rPr = r.get_or_add_rPr()
                            # You can add specific run formatting here if needed
                            
        except Exception as e:
            print(f"Warning while processing footnotes: {str(e)}")

    def _post_process_document(self, doc_path: Path, document_title: str) -> None:
        """Post-process the Word document with all necessary formatting"""
        doc = Document(doc_path)
        
        # Apply basic styles
        self._apply_global_styles(doc)
        self._process_paragraphs(doc)
        self._process_tables(doc)
        
        # Process footnotes
        self._process_footnotes(doc)
        
        # Setup footers
        self._setup_footers(doc)
        
        # Force title styling
        if doc.paragraphs:
            title_para = doc.paragraphs[0]
            title_para.clear()
            title_para.style = doc.styles['Title']
            run = title_para.add_run(document_title)
            run.font.name = self.config.font_name
            run.font.size = Pt(24)
            run.font.bold = True
        
        # Save the changes
        doc.save(doc_path)

    def _insert_images(self, doc_path: Path, image_refs: List[dict], work_dir: Path) -> None:
        """Insert images into the Word document"""
        doc = Document(doc_path)
        img_dir = work_dir / "img"
        
        for para in doc.paragraphs:
            if '[IMAGE_PLACEHOLDER]' in para.text and image_refs:
                img_ref = image_refs.pop(0)
                image_path = img_dir / img_ref['path']
                
                print(f"Processing image: {image_path}")
                
                if image_path.exists():
                    para.clear()
                    run = para.add_run()
                    try:
                        picture = run.add_picture(str(image_path))
                        
                        # Set image size
                        max_width = Inches(6)
                        if picture.width > max_width:
                            aspect_ratio = picture.height / picture.width
                            picture.width = max_width
                            picture.height = int(max_width * aspect_ratio)
                        
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"Added image: {img_ref['path']}")
                        
                    except Exception as e:
                        print(f"Error adding image {img_ref['path']}: {str(e)}")
                        para.text = f"[Image: {img_ref['alt_text']}]"
                else:
                    print(f"Image not found: {image_path}")
                    para.text = f"[Image not found: {img_ref['alt_text']}]"
        
        doc.save(doc_path)

    def _apply_global_styles(self, doc: Document) -> None:
        """Apply global styles to the document"""
        # Default styles
        styles_to_configure = [
            'Title', 'Normal', 'Body Text', 'Default Paragraph Font'
        ]
        
        for style_name in styles_to_configure:
            try:
                style = doc.styles[style_name]
                style.font.name = self.config.font_name
                if style_name == 'Title':
                    style.font.size = Pt(24)
                    style.font.bold = True
                else:
                    style.font.size = Pt(self.config.base_font_size)
            except KeyError:
                print(f"Style '{style_name}' not found")
        
        # Configure sections
        for section in doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            if self.config.paper_size == PaperSize.LEGAL:
                section.page_height = Inches(14)
            elif self.config.paper_size == PaperSize.A4:
                section.page_height = Inches(11.69)
                section.page_width = Inches(8.27)
            
            # Set margins
            section.top_margin = Cm(self.config.margins[0])
            section.right_margin = Cm(self.config.margins[1])
            section.bottom_margin = Cm(self.config.margins[2])
            section.left_margin = Cm(self.config.margins[3])

    def _process_paragraphs(self, doc: Document) -> None:
        """Process all paragraphs"""
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                self._apply_heading_style(para)
            else:
                for run in para.runs:
                    run.font.name = self.config.font_name
                    run.font.size = Pt(self.config.base_font_size)

    def _apply_heading_style(self, para) -> None:
        """Apply heading styles"""
        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
        color = self.config.heading_colors.get(level, (0, 0, 0))
        
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        
        for run in para.runs:
            run.font.name = self.config.font_name
            run.font.bold = True
            run.font.color.rgb = RGBColor(*color)

    def _process_tables(self, doc: Document) -> None:
        """Process all tables"""
        for table in doc.tables:
            # Process header row
            if table.rows:
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
            
            # Process all cells
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = self.config.font_name
                            run.font.size = Pt(self.config.base_font_size)


def create_report_example():
    """Example usage of the converter"""
    try:
        config = DocumentConfig.create_report_style(
            author="John Doe",
            date="2024-01-12",
            generate_toc=False,
            paper_size=PaperSize.LETTER,
            font_name="Arial",
            base_font_size=12,
            heading_colors={
                1: (0, 0, 0),
                2: (0, 0, 0),
                3: (0, 0, 0)
            },
            footer_text={
                "odd": "My Company",
                "even": "Confidential Report"
            }
        )
        
        # Chemin de travail absolu
        work_dir = "/Users/alain/Downloads"
        
        # Vérification du répertoire des images
        img_dir = Path(work_dir) / "img"
        if not img_dir.exists():
            print(f"Création du répertoire images: {img_dir}")
            img_dir.mkdir(parents=True)
        
        # Vérification de l'image
        image_path = img_dir / "markdown.png"
        if image_path.exists():
            print(f"Image trouvée: {image_path}")
            print(f"Taille de l'image: {os.path.getsize(image_path)} bytes")
        else:
            print(f"Image non trouvée: {image_path}")
        
        # Création et utilisation du convertisseur
        converter = MarkdownToDocxConverter(config)
        
        # Conversion avec messages de debug
        print(f"Démarrage de la conversion...")
        converter.convert(
            input_file="input.md",
            output_file="output3.docx",
            working_dir=work_dir
        )
        print(f"Conversion terminée")
        
    except Exception as e:
        print(f"Erreur lors de la conversion: {str(e)}")
        import traceback
        print(traceback.format_exc())

if __name__ == "__main__":
    create_report_example()
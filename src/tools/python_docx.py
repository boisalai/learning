"""
Create and modify Word documents from Markdown files using Python.

Requirements:
    pip install python-docx markdown beautifulsoup4

References:
    https://github.com/python-openxml/python-docx
"""
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class MarkdownToWordConverter:
    def __init__(self):
        self.document = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configure default document styles"""
        # Style du document
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Styles des titres
        for level in range(1, 4):
            style_name = f'Heading {level}'
            style = self.document.styles[style_name]
            font = style.font
            font.name = 'Calibri'
            font.bold = True
            if level == 1:
                font.size = Pt(16)
            elif level == 2:
                font.size = Pt(14)
            else:
                font.size = Pt(12)
                
    def add_headers_footers(self, header_text="", footer_text=""):
        """
        Add headers and footers to the document
        
        Args:
            header_text (str): Text to display in the header
            footer_text (str): Text to display in the footer
        """
        section = self.document.sections[0]
        
        # En-tête
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Pied de page
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajout numérotation des pages
        footer_para.add_run()
        self._add_page_number(footer_para)
        
    def _add_page_number(self, paragraph):
        """Add page numbering to the footer"""
        page_num = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        page_num._r.append(fldChar1)
        page_num._r.append(instrText)
        page_num._r.append(fldChar2)
        
    def convert_markdown(self, markdown_text):
        """
        Convert Markdown text to Word document format
        
        Args:
            markdown_text (str): Raw markdown text to convert
        """
        html = markdown.markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li', 'code', 'pre', 'blockquote']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                paragraph = self.document.add_paragraph(element.text)
                paragraph.style = self.document.styles[f'Heading {level}']
                
            elif element.name == 'p':
                paragraph = self.document.add_paragraph(element.text)
                
            elif element.name in ['ul', 'ol']:
                self._process_list(element)
                
            elif element.name == 'code':
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run(element.text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
            elif element.name == 'blockquote':
                paragraph = self.document.add_paragraph()
                paragraph.style = 'Quote'
                run = paragraph.add_run(element.text)
                paragraph.paragraph_format.left_indent = Inches(0.5)
                
    def _process_list(self, list_element, level=0):
        """
        Process nested lists (bulleted or numbered)
        
        Args:
            list_element: BeautifulSoup list element to process
            level (int): Current nesting level (default: 0)
        """
        for item in list_element.find_all('li', recursive=False):
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(level * 0.25)
            
            if list_element.name == 'ul':
                paragraph.style = 'List Bullet'
            else:
                paragraph.style = 'List Number'
                
            paragraph.add_run(item.text)
            
            # Traitement récursif des sous-listes
            nested_lists = item.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self._process_list(nested_list, level + 1)
                
    def save(self, output_path):
        """
        Save the document to a Word file
        
        Args:
            output_path (str): Path where the Word document will be saved
        """
        self.document.save(output_path)

def convert_file(input_path, output_path, header_text="", footer_text=""):
    """
    Convertit un fichier Markdown en document Word
    
    Args:
        input_path (str): Chemin du fichier Markdown source
        output_path (str): Chemin du fichier Word de destination
        header_text (str): Texte à placer dans l'en-tête
        footer_text (str): Texte à placer dans le pied de page
    """
    # Lecture du fichier Markdown
    with open(input_path, 'r', encoding='utf-8') as file:
        markdown_text = file.read()
    
    # Conversion
    converter = MarkdownToWordConverter()
    converter.convert_markdown(markdown_text)
    
    # Ajout en-têtes et pieds de page
    if header_text or footer_text:
        converter.add_headers_footers(header_text, footer_text)
    
    # Sauvegarde
    converter.save(output_path)

# Exemple d'utilisation
if __name__ == "__main__":
    convert_file(
        "../../data/cargo_cult_science_feynman_1974.txt",
        "../../data/output.docx",
        header_text="Mon Document",
        footer_text="Confidentiel"
    )
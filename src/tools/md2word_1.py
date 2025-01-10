import os
import subprocess
import shutil
from pathlib import Path
from typing import Optional, Dict, List
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.table import _Cell

class MarkdownToDocx:
    def __init__(self):
        self._check_dependencies()
    
    def _process_image_paths(self, content: str, img_dir: Path) -> str:
        """
        Traite les chemins des images dans le contenu markdown
        - Vérifie l'existence des images dans le répertoire img/
        - Ajuste les chemins pour pointer vers le répertoire img/
        """
        import re
        
        # Motif pour détecter les images en markdown
        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        
        def process_image_match(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            
            # Obtenir juste le nom du fichier image
            img_name = Path(img_path).name
            
            # Vérifier si l'image existe dans le répertoire img/
            img_file = img_dir / img_name
            if not img_file.exists():
                print(f"Attention: Image non trouvée: {img_file}")
                return match.group(0)  # Garder le chemin original si l'image n'existe pas
                
            # Retourner le markdown avec le chemin ajusté
            return f'![{alt_text}](img/{img_name})'
            
        # Remplacer tous les chemins d'images
        return re.sub(img_pattern, process_image_match, content)
        
    def _check_dependencies(self) -> None:
        if not shutil.which('pandoc'):
            raise RuntimeError(
                "Pandoc n'est pas installé. "
                "Sur macOS, installez-le avec: brew install pandoc"
            )

    def _set_cell_borders(self, cell: _Cell):
        """Ajoute des bordures à une cellule"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        
        # Correction : utiliser 'is None' au lieu de la vérification booléenne
        if tcBorders is None:  # Modifié ici
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '4')
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), 'auto')
            tcBorders.append(border_el)

    def _remove_list_spacing(self, paragraph) -> None:
        """Supprime les espacements d'un paragraphe au niveau XML"""
        pPr = paragraph._p.get_or_add_pPr()
        
        # Supprimer l'espacement avant
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Interligne simple
        spacing.set(qn('w:lineRule'), 'auto')
        
        # Supprimer l'ancien élément spacing s'il existe
        old_spacing = pPr.find(qn('w:spacing'))
        if old_spacing is not None:
            pPr.remove(old_spacing)
        
        pPr.append(spacing)

    def _setup_footers(self, doc):
        """Configure des pieds de page distincts pour les pages paires et impaires."""
        for section in doc.sections:
            # Activer les options pour différents pieds de page
            section.different_first_page_header_footer = True  
            section.odd_and_even_pages_header_footer = True 

            # Pied de page des pages impaires (droites)
            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()

            # Ajouter l'espace de 12pt avant
            p_odd.paragraph_format.space_before = Pt(12)
            p_odd.paragraph_format.space_after = Pt(0)

            run_odd = p_odd.add_run("Texte pied de page impaire | ")
            run_odd.font.name = 'Arial'
            run_odd.font.size = Pt(10)
            self._add_page_number(p_odd)
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Pied de page des pages paires (gauches)
            # Utiliser even_page_footer au lieu de footer pour les pages paires
            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()

            # Ajouter l'espace de 12pt avant
            p_even.paragraph_format.space_before = Pt(12)
            p_even.paragraph_format.space_after = Pt(0)

            self._add_page_number(p_even)
            run_even = p_even.add_run(" | Texte pied de page paire")
            run_even.font.name = 'Arial' 
            run_even.font.size = Pt(10)
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _configure_section_properties(self, doc):
        """Configure les propriétés de section pour les en-têtes et pieds de page différents"""
        for section in doc.sections:
            # Accéder aux propriétés XML de la section
            sectPr = section._sectPr
            
            # Créer l'élément pour les pages paires/impaires différentes
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            
            # Activer les en-têtes/pieds de page différents pour pages paires/impaires
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

    def _process_footnotes(self, doc):
        """Configure les notes de bas de page en Arial 10pt sans espacement"""
        try:
            # Accéder aux notes de bas de page via le style intégré
            styles = doc.styles
            style = styles['Footnote Text']  # Style intégré pour les notes de bas de page
            
            # Configurer la police et la taille
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            
            # Supprimer les espacements avant/après
            style.paragraph_format.space_before = Pt(0)  # Ajoute 12pt avant
            style.paragraph_format.space_after = Pt(0)  # Garde 0pt après
            style.paragraph_format.line_spacing = 1.0  # Interligne simple
            
            # Configurer le style des caractères de référence de note de bas de page
            ref_style = styles['Footnote Reference']
            ref_style.font.name = 'Arial'
            ref_style.font.size = Pt(10)
            
            # Traiter chaque note de bas de page individuellement si nécessaire
            if hasattr(doc.part, '_footnotes_part') and doc.part._footnotes_part:
                footnotes = doc.part._footnotes_part.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                for footnote in footnotes:
                    # Force l'application du style et des espacements au niveau XML
                    pPr = footnote.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                    if pPr is None:
                        pPr = parse_xml(r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        footnote.insert(0, pPr)
                    
                    # Configurer l'espacement
                    spacing = parse_xml(
                        r'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        r'w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
                    )
                    
                    # Supprimer l'ancien espacement s'il existe
                    old_spacing = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                    if old_spacing is not None:
                        pPr.remove(old_spacing)
                    
                    pPr.append(spacing)

        except Exception as e:
            print(f"Avertissement lors du traitement des notes de bas de page : {str(e)}")


    def _adjust_heading_number_spacing(self, doc):
        """Ajuste l'espacement entre la numérotation et le texte des titres"""
        for style in doc.styles:
            if style.name.startswith('Heading'):
                if hasattr(style, '_element') and style._element.pPr is not None:
                    pPr = style._element.pPr
                    
                    # Configuration de la numérotation
                    numPr = pPr.get_or_add_numPr()
                    
                    # Configurer les tabs après la numérotation
                    tabs = pPr.get_or_add_tabs()
                    
                    # Supprimer tous les tabs existants
                    existing_tabs = tabs.findall(qn('w:tab'))
                    for tab in existing_tabs:
                        tabs.remove(tab)
                    
                    # Ajouter un nouveau tab pour le texte
                    tab = OxmlElement('w:tab')
                    tab.set(qn('w:val'), 'left')
                    # tab.set(qn('w:pos'), '360')  # 18pt
                    # tab.set(qn('w:pos'), '720')  # 36pt - augmenté de 360 à 720
                    tab.set(qn('w:pos'), '960')  # 48pt
                    tabs.append(tab)
                    
                    # Réinitialiser toutes les indentations
                    ind = pPr.get_or_add_ind()
                    ind.set(qn('w:left'), '0')
                    ind.set(qn('w:firstLine'), '0')
                    ind.set(qn('w:hanging'), '0')


    def _adjust_toc_styles(self, doc):
        """Supprime le gras des paragraphes de la table des matières"""
        in_toc = False
        
        for paragraph in doc.paragraphs:
            # Détecter le début de la table des matières
            if paragraph.text == "Table des matières" or paragraph.text == "Table of Contents":
                in_toc = True
                continue
                
            # Détecter la fin de la table des matières (quand on trouve un titre de niveau 1)
            if in_toc and paragraph.style.name.startswith('Heading 1'):
                in_toc = False
                continue
                
            # Modifier les paragraphes de la table des matières
            if in_toc:
                # Supprimer le gras au niveau des runs
                for run in paragraph.runs:
                    run.font.bold = False
                    
                    # Modification au niveau XML pour s'assurer que le gras est désactivé
                    if hasattr(run, '_element') and run._element.rPr is not None:
                        rPr = run._element.rPr
                        # Supprimer tout attribut bold existant
                        bold_elements = rPr.findall(qn('w:b'))
                        for bold in bold_elements:
                            rPr.remove(bold)
                        # Ajouter un attribut bold explicitement désactivé
                        bold = OxmlElement('w:b')
                        bold.set(qn('w:val'), '0')
                        rPr.append(bold)


    def _post_process_document(self, doc_path: Path) -> None:
        """Post-traitement du document Word"""
        doc = Document(doc_path)

        # Ajouter l'ajustement des styles TOC
        self._adjust_toc_styles(doc)

        # Ajouter l'ajustement des espacements de numérotation
        self._adjust_heading_number_spacing(doc)

        # Configuration des propriétés de section
        self._configure_section_properties(doc)
        
        # Configuration des pieds de page
        self._setup_footers(doc)
        
        # Configuration des notes de bas de page
        self._process_footnotes(doc)
            
        # Configuration des marges et des sections
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Traitement des paragraphes
        in_list = False
        for para in doc.paragraphs:
            # Détecter si nous sommes dans une liste
            if ('•' in para.text or 
                para.text.strip().startswith('o') or 
                para.text.strip().startswith('§') or
                any(str(i)+'.' in para.text for i in range(10))):
                
                self._remove_list_spacing(para)
                in_list = True
            else:
                in_list = False
            
            # Appliquer les styles
            if para.style.name.startswith('Heading'):
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after = Pt(12)
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(37, 150, 190)
            
            elif 'CodeBlock' in para.style.name:
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                para.paragraph_format.left_indent = Cm(1)
            
            elif "Logo Markdown" in para.text:
                image_path = doc_path.parent / "markdown.png"
                if image_path.exists():
                    para.text = ""
                    run = para.add_run()
                    run.add_picture(str(image_path), width=Inches(3))
    
        # Traitement des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = para.text.strip()
                        para.paragraph_format.space_before = 0
                        para.paragraph_format.space_after = 0
                    self._set_cell_borders(cell)

            # Style de base
            font = para.style.font
            font.name = 'Arial'
            font.size = Pt(12)
            
            # Styles spécifiques
            if 'CodeBlock' in para.style.name:
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                para.paragraph_format.left_indent = Cm(1)
            elif para.style.name.startswith('Heading'):
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(37, 150, 190)

            # Gestion des images
            if "Logo Markdown" in para.text:
                image_path = downloads_dir / "markdown.png"
                if image_path.exists():
                    para.text = ""  # Effacer le texte existant
                    run = para.add_run()
                    run.add_picture(str(image_path), width=Inches(3))
                else:
                    print(f"Image non trouvée: {image_path}")

        # Traitement des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = para.text.strip()
                        para.paragraph_format.space_before = 0
                        para.paragraph_format.space_after = 0
                    self._set_cell_borders(cell)

        # Sauvegarder les modifications
        doc.save(doc_path)

    def _add_page_number(self, paragraph):
        """Ajoute un numéro de page au paragraphe"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def convert(self, 
                input_file: str, 
                output_file: str,
                metadata: Optional[Dict] = None,
                extra_args: Optional[List[str]] = None) -> None:
        """Convertit un fichier Markdown en Word"""
        input_path = Path(input_file).resolve()
        output_path = Path(output_file).resolve()
        
        if not input_path.exists():
            raise FileNotFoundError(f"Le fichier d'entrée n'existe pas: {input_path}")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Commande pandoc de base
        cmd = [
            'pandoc',
            str(input_path),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999'
        ]
        
        # Ajouter les métadonnées
        if metadata:
            for key, value in metadata.items():
                cmd.extend(['-M', f'{key}={value}'])
        
        # Ajouter les arguments supplémentaires
        if extra_args:
            cmd.extend(extra_args)
        
        try:
            # Conversion avec pandoc
            result = subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                text=True
            )
            
            # Post-traitement du document
            self._post_process_document(output_path)
            
            print(f"Conversion réussie! Fichier sauvegardé: {output_path}")
            print("\nNote pour les images:")
            print("- Placez les images dans le même dossier que votre fichier markdown")
            print("- Utilisez des chemins relatifs: ![titre](image.png)")
            print("- Formats supportés: PNG, JPEG, GIF, BMP")
            
        except subprocess.CalledProcessError as e:
            print(f"Erreur lors de la conversion: {e.stderr}")
            raise

if __name__ == "__main__":
    try:
        # Définir les chemins
        downloads_dir = Path("/Users/alain/Downloads")
        input_path = downloads_dir / "input.md"
        output_path = downloads_dir / "output.docx"
        
        # Créer et utiliser le convertisseur
        converter = MarkdownToDocx()
        
        # Définir les métadonnées
        metadata = {
            "title": "Titre principal du document",
            "author": "auteur",
            "date": "2024-01-09"
        }
        
        # Options supplémentaires
        extra_args = [
            '--toc',
            '--toc-depth=2',
            '--number-sections',
            '--variable', 'toc-title=Table des matières' 
        ]
        
        # Convertir le fichier
        converter.convert(
            input_file=input_path,
            output_file=output_path,
            metadata=metadata,
            extra_args=extra_args
        )
        
    except Exception as e:
        print(f"Erreur: {str(e)}")
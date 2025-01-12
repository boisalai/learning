from typing import Optional, Dict, List, Tuple
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import os
import subprocess
import shutil


class PaperSize(Enum):
    LETTER = "letter"  # 8.5 x 11 inches (215.9 x 279.4 mm)
    LEGAL = "legal"    # 8.5 x 14 inches (215.9 x 355.6 mm)
    A4 = "a4"         # 8.27 x 11.69 inches (210 x 297 mm)

class DocumentStyle(Enum):
    REPORT = "report"
    NOTE = "note"
    LETTER = "letter"
    MEMO = "memo"

class DocumentConfig:
    def __init__(
        self,
        style: DocumentStyle = DocumentStyle.REPORT,
        paper_size: PaperSize = PaperSize.LETTER,
        author: str = "",
        date: str = "",
        heading_colors: Dict[int, Tuple[int, int, int]] = None,
        footer_text: Dict[str, str] = None,
        font_name: str = "Arial",
        base_font_size: int = 12,
        margins: Tuple[float, float, float, float] = (2, 2, 2, 2),  # top, right, bottom, left in cm
        line_spacing: float = 1.0,
        generate_toc: bool = True
    ):
        self.style = style
        self.paper_size = paper_size
        self.author = author
        self.date = date
        self.heading_colors = heading_colors or {
            1: (37, 150, 190),  # Default blue color for all headings
            2: (37, 150, 190),
            3: (37, 150, 190)
        }
        self.footer_text = footer_text or {
            "odd": "Page",
            "even": "Page"
        }
        self.font_name = font_name
        self.base_font_size = base_font_size
        self.margins = margins
        self.line_spacing = line_spacing
        self.generate_toc = generate_toc

    @classmethod
    def create_report_style(cls, **kwargs):
        """Preconfigured style for reports"""
        default_config = {
            'style': DocumentStyle.REPORT,
            'paper_size': PaperSize.LETTER,
            'heading_colors': {
                1: (37, 150, 190),
                2: (37, 150, 190),
                3: (37, 150, 190)
            },
            'footer_text': {
                "odd": "Right text | Page",
                "even": "Page | Left text"
            }
        }
        default_config.update(kwargs)
        return cls(**default_config)

    @classmethod
    def create_note_style(cls, **kwargs):
        """Preconfigured style for internal notes"""
        default_config = {
            'style': DocumentStyle.NOTE,
            'paper_size': PaperSize.LEGAL,
            'heading_colors': {
                1: (70, 70, 70),
                2: (100, 100, 100),
                3: (130, 130, 130)
            },
            'footer_text': {
                "odd": "Internal Note | Page",
                "even": "Page | Internal Note"
            },
            'margins': (1.5, 1.5, 1.5, 1.5)
        }
        default_config.update(kwargs)
        return cls(**default_config)

class MarkdownToDocxConverter:
    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if Pandoc is installed"""
        if not shutil.which('pandoc'):
            raise RuntimeError(
                "Pandoc is not installed. "
                "On macOS, install it with: brew install pandoc"
            )

    def _extract_title_from_markdown(self, content: str) -> str:
        """
        Extract the main title (h1) from markdown content
        
        Args:
            content: The markdown content as string
            
        Returns:
            str: The main title without the '# ' prefix, or 'Untitled Document' if no h1 found
        """
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
        return "Untitled Document"

    def convert(self, 
                input_file: str,
                output_file: str,
                working_dir: Optional[str] = None,
                extra_args: Optional[List[str]] = None) -> None:
        """
        Convert a Markdown file to Word with the specified configuration
        
        Args:
            input_file: Name of the input markdown file
            output_file: Name of the output Word file
            working_dir: Working directory (containing the markdown and where the docx will be created)
            extra_args: Additional arguments for pandoc
        """
        # Set the working directory
        work_dir = Path(working_dir).resolve() if working_dir else Path.cwd()
        input_path = work_dir / input_file
        output_path = work_dir / output_file
        
        # Check if the input file exists
        if not input_path.exists():
            raise FileNotFoundError(f"Input file does not exist: {input_path}")
        
        # Read the markdown file content to check for images
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Extract title from markdown content
        document_title = self._extract_title_from_markdown(content)

        # Adjust image paths in the markdown content
        content = self._process_image_paths(content, img_dir)
        
        # Create a temporary markdown file with adjusted paths
        temp_md = work_dir / f"temp_{input_path.name}"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Pandoc command
        cmd = [
            'pandoc',
            str(temp_md),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999',
            '--resource-path', str(img_dir),
            '-M', f'title={document_title}',
            '-M', f'author={self.config.author}',
            '-M', f'date={self.config.date}'
        ]
        
        # Add table of contents if enabled
        if self.config.generate_toc:
            cmd.extend(['--toc', '--number-sections'])

        if extra_args:
            cmd.extend(extra_args)
        
        try:
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            self._post_process_document(output_path, document_title)
            print(f"Conversion successful! File saved: {output_path}")
            
        except subprocess.CalledProcessError as e:
            print(f"Error during conversion: {e.stderr}")
            raise
            
        finally:
            # Clean up the temporary file
            if temp_md.exists():
                temp_md.unlink()

    def _process_image_paths(self, content: str, img_dir: Path) -> str:
        """Process image paths in the markdown content"""
        import re
        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        
        print(f"Looking for images in: {img_dir}") 

        def process_image_match(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            img_name = Path(img_path).name
            img_file = img_dir / img_name
            
            if not img_file.exists():
                print(f"Warning: Image not found: {img_file}")
                return match.group(0)
                
            print(f"Processed image path: {img_file}")  # Ajout d'un log pour débogage
            return f'![{alt_text}](img/{img_name})'
            
        return re.sub(img_pattern, process_image_match, content)

    def _post_process_document(self, doc_path: Path, document_title: str) -> None:
        """
        Post-process the Word document to apply all styles and formatting.
        """
        doc = Document(doc_path)

        # Apply global styles and configurations
        self._apply_global_styles(doc)

        # Process special paragraphs (lists, headings, code, etc.)
        self._process_paragraphs(doc, doc_path)
        
        # Process tables
        self._process_tables(doc)
        
        # Force the title style (first paragraph)
        title_para = doc.paragraphs[0]
        title_para.clear()
        title_para.style = doc.styles['Title']
        run = title_para.add_run(document_title)
        run.font.name = self.config.font_name
        run.font.size = Pt(24)
        run.font.bold = True

        # Save the changes
        doc.save(doc_path)

    def _apply_global_styles(self, doc: Document) -> None:
        """Apply global styles and configurations to the document"""
        try:
            title_style = doc.styles['Title']
            title_style.font.name = self.config.font_name
            title_style.font.size = Pt(24)
            title_style.font.bold = True
        except KeyError:
            print("Style 'Title' not found")

        style_normal = doc.styles['Normal']
        style_normal.font.name = self.config.font_name
        style_normal.font.size = Pt(self.config.base_font_size)
        style_normal.paragraph_format.line_spacing = self.config.line_spacing

        try:
            style_default = doc.styles['Default Paragraph Font']
            style_default.font.name = self.config.font_name
            style_default.font.size = Pt(self.config.base_font_size)
        except KeyError:
            print("Style 'Default Paragraph Font' not found")

        for style_name in ['Body Text', 'Normal', 'Plain Text']:
            try:
                style = doc.styles[style_name]
                style.font.name = self.config.font_name
                style.font.size = Pt(self.config.base_font_size)
            except KeyError:
                continue

        self._adjust_toc_styles(doc)
        self._adjust_heading_number_spacing(doc)
        self._configure_section_properties(doc)
        self._setup_footers(doc)
        self._process_footnotes(doc)
        
        for section in doc.sections:
            section.top_margin = Cm(self.config.margins[0])
            section.bottom_margin = Cm(self.config.margins[1])
            section.left_margin = Cm(self.config.margins[2])
            section.right_margin = Cm(self.config.margins[3])

    def _process_image_paragraph(self, para, doc_path: Path) -> None:
        """Process a paragraph containing an image reference"""
        import re
        
        print(">>>" + doc_path)

        # Extract image info using regex
        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        matches = re.findall(img_pattern, para.text)
        
        if matches:
            for alt_text, img_path in matches:
                # Clean up image path
                img_path = img_path.strip()
                if img_path.startswith('img/'):
                    img_path = img_path[4:]  # Remove 'img/' prefix
                
                # Construct full image path
                image_path = doc_path.parent / "img" / img_path
                
                print(f"Processing image at: {image_path}")
                
                if image_path.exists():
                    # Clear existing text
                    para.clear()
                    
                    # Add image
                    run = para.add_run()
                    try:
                        picture = run.add_picture(str(image_path))
                        
                        # Set image size (optional, adjust as needed)
                        max_width = Inches(6)  # Maximum width of 6 inches
                        if picture.width > max_width:
                            aspect_ratio = picture.height / picture.width
                            picture.width = max_width
                            picture.height = int(max_width * aspect_ratio)
                        
                        # Center the paragraph
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        print(f"Successfully added image: {img_path}")
                    except Exception as e:
                        print(f"Error adding image {img_path}: {str(e)}")
                        para.text = f"[Image: {alt_text}]"
                else:
                    print(f"Image not found: {image_path}")
                    para.text = f"[Image not found: {alt_text}]"

    def _process_paragraphs(self, doc: Document, doc_path: Path) -> None:
        """Process all paragraphs in the document"""
        for para in doc.paragraphs:
            # Detect if paragraph contains an image reference
            if '![' in para.text and '](' in para.text:
                self._process_image_paragraph(para, doc_path)
            else:
                if not para.style.name.startswith('Heading'):
                    for run in para.runs:
                        run.font.name = self.config.font_name
                        run.font.size = Pt(self.config.base_font_size)

                if self._is_list_paragraph(para.text):
                    self._remove_list_spacing(para)
                
                if para.style.name.startswith('Heading'):
                    self._apply_heading_style(para)
                elif 'CodeBlock' in para.style.name:
                    self._apply_code_style(para)

    def _is_list_paragraph(self, text: str) -> bool:
        """Determine if a paragraph is a list"""
        return ('•' in text or 
                text.strip().startswith('o') or 
                text.strip().startswith('§') or
                any(str(i)+'.' in text for i in range(10)))

    def _apply_heading_style(self, para) -> None:
        """Apply style to headings"""
        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        for run in para.runs:
            run.font.name = self.config.font_name
            run.font.bold = True
            run.font.color.rgb = RGBColor(*self.config.heading_colors.get(level, (37, 150, 190)))

    def _apply_code_style(self, para) -> None:
        """Apply style to code blocks"""
        for run in para.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        para.paragraph_format.left_indent = Cm(1)

    def _process_logo_image(self, para, doc_path: Path) -> None:
        """Process the Markdown logo image"""
        image_path = doc_path.parent / "img" / "markdown.png"
        print(image_path)
        if image_path.exists():
            para.text = ""
            run = para.add_run()
            run.add_picture(str(image_path), width=Inches(3))
        else:
            print(f"Image not found: {image_path}")

    def _process_tables(self, doc: Document) -> None:
        """Process all tables in the document"""
        for table in doc.tables:
            self._process_table_header(table)
            for i, row in enumerate(table.rows):
                if i > 0:  
                    for cell in row.cells:
                        self._process_table_cell(cell, is_header=False)

    def _process_table_header(self, table) -> None:
        """Process the first row of the table (header)"""
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                self._process_table_cell(cell, is_header=True)

    def _process_table_cell(self, cell: _Cell, is_header: bool = False) -> None:
        """Process a table cell"""
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            table_font_size = self.config.base_font_size - 1
            for run in para.runs:
                run.font.name = self.config.font_name
                run.font.size = Pt(table_font_size)
                if is_header:
                    run.font.bold = True
        
            if not para.text.strip():
                run = para.add_run(" ")
                run.font.name = self.config.font_name
                run.font.size = Pt(table_font_size)
                if is_header:
                    run.font.bold = True
                    
        self._set_cell_borders(cell)

    def _set_cell_borders(self, cell: _Cell) -> None:
        """Add borders to a cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '4')
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), 'auto')
            tcBorders.append(border_el)

    def _remove_list_spacing(self, paragraph) -> None:
        """Remove spacing from a paragraph at the XML level"""
        pPr = paragraph._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        old_spacing = pPr.find(qn('w:spacing'))
        if old_spacing is not None:
            pPr.remove(old_spacing)
        pPr.append(spacing)

    def _configure_section_properties(self, doc) -> None:
        """Configure section properties, including paper size"""
        for section in doc.sections:
            # Configure paper size
            if self.config.paper_size == PaperSize.LETTER:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
            elif self.config.paper_size == PaperSize.LEGAL:
                section.page_width = Inches(8.5)
                section.page_height = Inches(14)
            else:  # A4
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)

            # Configure margins
            section.top_margin = Cm(self.config.margins[0])
            section.right_margin = Cm(self.config.margins[1])
            section.bottom_margin = Cm(self.config.margins[2])
            section.left_margin = Cm(self.config.margins[3])

            # Configure properties for headers and footers
            sectPr = section._sectPr
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

    def _setup_footers(self, doc) -> None:
        """Configure footers with custom text"""
        for section in doc.sections:
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True

            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()
            p_odd.paragraph_format.space_before = Pt(12)
            run_odd = p_odd.add_run(f"{self.config.footer_text['odd']} | ")
            run_odd.font.name = self.config.font_name
            run_odd.font.size = Pt(10)
            self._add_page_number(p_odd)
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()
            p_even.paragraph_format.space_before = Pt(12)
            self._add_page_number(p_even)
            run_even = p_even.add_run(f" | {self.config.footer_text['even']}")
            run_even.font.name = self.config.font_name
            run_even.font.size = Pt(10)
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_page_number(self, paragraph) -> None:
        """Add a page number to the paragraph"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _process_footnotes(self, doc) -> None:
        """Configure footnotes"""
        try:
            styles = doc.styles
            style = styles['Footnote Text']
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            ref_style = styles['Footnote Reference']
            ref_style.font.name = 'Arial'
            ref_style.font.size = Pt(10)
            
            if hasattr(doc.part, '_footnotes_part') and doc.part._footnotes_part:
                footnotes = doc.part._footnotes_part.element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
                )
                for footnote in footnotes:
                    pPr = footnote.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr'
                    )
                    if pPr is None:
                        pPr = parse_xml(
                            r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        )
                        footnote.insert(0, pPr)
                    spacing = parse_xml(
                        r'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        r'w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
                    )
                    old_spacing = pPr.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing'
                    )
                    if old_spacing is not None:
                        pPr.remove(old_spacing)
                    pPr.append(spacing)

        except Exception as e:
            print(f"Warning while processing footnotes: {str(e)}")

    def _adjust_toc_styles(self, doc) -> None:
        """Remove bold from table of contents paragraphs"""
        in_toc = False
        for paragraph in doc.paragraphs:
            if paragraph.text in ["Table des matières", "Table of Contents"]:
                in_toc = True
                continue
            if in_toc and paragraph.style.name.startswith('Heading 1'):
                in_toc = False
                continue
            if in_toc:
                for run in paragraph.runs:
                    run.font.bold = False
                    if hasattr(run, '_element') and run._element.rPr is not None:
                        rPr = run._element.rPr
                        for bold in rPr.findall(qn('w:b')):
                            rPr.remove(bold)
                        bold = OxmlElement('w:b')
                        bold.set(qn('w:val'), '0')
                        rPr.append(bold)

    def _adjust_heading_number_spacing(self, doc) -> None:
        """Adjust spacing between numbering and heading text"""
        for style in doc.styles:
            if style.name.startswith('Heading'):
                if hasattr(style, '_element') and style._element.pPr is not None:
                    pPr = style._element.pPr
                    numPr = pPr.get_or_add_numPr()
                    tabs = pPr.get_or_add_tabs()
                    for tab in tabs.findall(qn('w:tab')):
                        tabs.remove(tab)
                    tab = OxmlElement('w:tab')
                    tab.set(qn('w:val'), 'left')
                    tab.set(qn('w:pos'), '960')
                    tabs.append(tab)
                    ind = pPr.get_or_add_ind()
                    ind.set(qn('w:left'), '0')
                    ind.set(qn('w:firstLine'), '0')
                    ind.set(qn('w:hanging'), '0')

def create_note_example():
    try:
        converter = MarkdownToDocxConverter(
            config=DocumentConfig.create_report_style(
                author="John Doe",
                date="2024-01-10"
            )
        )
        converter.convert(
            input_file="input.md",
            output_file="output1.docx",
            working_dir="/Users/alain/Downloads",
            extra_args=['--toc', '--number-sections']
        )
    except Exception as e:
        print(f"Error: {str(e)}")

def create_report_example():
    try:
        config = DocumentConfig.create_report_style(
            author="John Doe",
            date="2024-01-10",
            generate_toc=False,
            paper_size=PaperSize.LEGAL,
            font_name="Arial", 
            base_font_size=12,  
            heading_colors={
                1: (0, 0, 0),
                2: (0, 0, 0),
                3: (0, 0, 0)
            },
            footer_text={
                "odd": "My Company",
                "even": "Confidential Report"
            }
        )
        converter = MarkdownToDocxConverter(config)
        converter.convert(
            input_file="input.md",
            output_file="output3.docx",
            working_dir="/Users/alain/Downloads"        
        )
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    # create_note_example()
    create_report_example()